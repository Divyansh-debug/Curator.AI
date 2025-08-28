import os
import requests
import logging
import google.generativeai as genai
from bs4 import BeautifulSoup
from fastapi import APIRouter, File, UploadFile, HTTPException
from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document as LangchainDocument

# Configure the Gemini API key from environment variables
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Set up logging
logging.basicConfig(level=logging.INFO)

router = APIRouter()

# Define a function to process documents and create a FAISS index
def process_and_index_documents(texts: list[str], index_name: str = "faiss_index"):
    """
    Processes a list of text documents, creates embeddings using Gemini, and stores them in a FAISS index.
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    docs = text_splitter.create_documents(texts)
    
    if not docs:
        logging.error("No documents were created. Check the input text.")
        return

    logging.info(f"Creating a FAISS index with {len(docs)} documents.")
    
    # Generate embeddings using Gemini API
    texts_to_embed = [doc.page_content for doc in docs]
    embeddings_response = genai.embed_content(
        model="models/embedding-001",
        content=texts_to_embed,
        task_type="retrieval_document"
    )
    
    # Create the FAISS index with the embeddings and document content
    embeddings_list = embeddings_response['embedding']
    db = FAISS.from_embeddings(
        text_embeddings=list(zip(texts_to_embed, embeddings_list)),
        embedding=None, # Not needed as we already have embeddings
        metadatas=[doc.metadata for doc in docs]
    )
    
    db.save_local(index_name)
    logging.info(f"FAISS index saved to '{index_name}'.")

# File Upload Endpoint
@router.post("/ingest-documents/")
async def ingest_documents(files: list[UploadFile] = File(...)):
    """
    Ingests PDF and Word documents, extracts text, and creates a FAISS index.
    """
    all_text = []
    
    for file in files:
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        try:
            if file_extension == ".pdf":
                reader = PdfReader(file.file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                all_text.append(text)
                
            elif file_extension in [".docx", ".doc"]:
                doc = Document(file.file)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                all_text.append(text)
                
            else:
                logging.warning(f"Skipping unsupported file type: {file.filename}")
                continue
                
        except Exception as e:
            logging.error(f"Error processing file {file.filename}: {e}")
            raise HTTPException(status_code=500, detail=f"Error processing file {file.filename}")
            
    if all_text:
        process_and_index_documents(all_text)
        return {"status": "success", "message": f"Successfully ingested {len(files)} documents."}
    else:
        return {"status": "error", "message": "No valid documents were processed."}

# Web Scraping Endpoint
@router.post("/scrape-website/")
async def scrape_website(url: str):
    """
    Scrapes a given URL, extracts text, and updates the FAISS index.
    """
    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="URL must start with http:// or https://")
    
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        
        body_text = soup.body.get_text(separator="\n", strip=True)
        
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        docs = text_splitter.create_documents([body_text])
        texts_to_embed = [doc.page_content for doc in docs]
        
        embeddings_response = genai.embed_content(
            model="models/embedding-001",
            content=texts_to_embed,
            task_type="retrieval_document"
        )
        embeddings_list = embeddings_response['embedding']
        
        # Load the existing index or create a new one
        if os.path.exists("faiss_index"):
            logging.info("Loading existing FAISS index...")
            
            # Since LangChain's FAISS class has a custom deserialization that needs the embedding function,
            # it's simpler to directly interact with the faiss library for updates.
            # For this simple example, let's just re-create the index. For a production app,
            # you'd implement logic to add to the existing index.
            
            # Simple approach: append new docs and rebuild index
            old_db = FAISS.load_local("faiss_index", embeddings=None, allow_dangerous_deserialization=True)
            old_docs = [LangchainDocument(page_content=text) for text in old_db.docstore._dict.values()]
            
            new_docs_texts = [doc.page_content for doc in docs]
            all_docs_texts = [d.page_content for d in old_docs] + new_docs_texts
            
            process_and_index_documents(all_docs_texts)
        else:
            logging.info("No existing FAISS index found. Creating a new one.")
            process_and_index_documents([body_text])

        return {"status": "success", "message": f"Successfully scraped and ingested data from {url}."}
        
    except requests.exceptions.RequestException as e:
        logging.error(f"Error scraping URL {url}: {e}")
        raise HTTPException(status_code=500, detail=f"Error scraping URL: {e}")
    except Exception as e:
        logging.error(f"An unexpected error occurred: {e}")
        raise HTTPException(status_code=500, detail="An unexpected server error occurred.")
